"""
Step 5: Word Document Loaders for Ingestion and Parsing

This module provides Word document loaders for .docx files.
Supports multiple loading strategies:
- python-docx (direct parsing)
- Docx2txtLoader (langchain, simple text extraction)
- UnstructuredWordDocumentLoader (langchain, structured element extraction)
"""
import os
import tempfile
from typing import List, Optional
from pathlib import Path
from langchain_core.documents import Document

from app.core.logging_config import get_logger

logger = get_logger(__name__)


class WordDocumentLoader:
    """
    Loader for Word (.docx) files
    
    Supports:
    - Text extraction from paragraphs
    - Table extraction
    - Metadata extraction (author, title, etc.)
    - Section-based chunking
    """
    
    def __init__(
        self,
        extract_tables: bool = True,
        extract_metadata: bool = True,
        loader_type: str = "python-docx"
    ):
        """
        Initialize the Word document loader
        
        Args:
            extract_tables: Whether to extract text from tables (python-docx only)
            extract_metadata: Whether to extract document metadata
            loader_type: Type of loader to use ('python-docx', 'docx2txt', 'unstructured')
        """
        self.extract_tables = extract_tables
        self.extract_metadata = extract_metadata
        self.loader_type = loader_type.lower()
        
        if self.loader_type not in ['python-docx', 'docx2txt', 'unstructured']:
            logger.warning(f"Unknown loader_type '{loader_type}', defaulting to 'python-docx'")
            self.loader_type = 'python-docx'
    
    def _extract_text_from_paragraphs(self, doc) -> str:
        """Extract text from all paragraphs in the document"""
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Skip empty paragraphs
                paragraphs.append(text)
        return "\n\n".join(paragraphs)
    
    def _extract_text_from_tables(self, doc) -> str:
        """Extract text from tables in the document"""
        if not self.extract_tables:
            return ""
        
        table_texts = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                table_texts.append("\n".join(table_rows))
        
        return "\n\n--- Table ---\n\n".join(table_texts) if table_texts else ""
    
    def _extract_metadata(self, doc) -> dict:
        """Extract metadata from Word document"""
        if not self.extract_metadata:
            return {}
        
        core_props = doc.core_properties
        
        metadata = {}
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.subject:
            metadata["subject"] = core_props.subject
        if core_props.created:
            metadata["created"] = core_props.created.isoformat() if core_props.created else None
        if core_props.modified:
            metadata["modified"] = core_props.modified.isoformat() if core_props.modified else None
        if core_props.comments:
            metadata["comments"] = core_props.comments
        
        return metadata
    
    def load_file(self, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """
        Load a Word (.docx) file and return Document objects
        
        Args:
            file_path: Path to the .docx file
            metadata: Optional metadata to add to documents
            
        Returns:
            List of Document objects (typically one document, but can be split by sections/elements)
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        logger.debug(f"Loading Word file | Path: {file_path} | Loader: {self.loader_type}")
        
        # Route to appropriate loader based on loader_type
        if self.loader_type == 'docx2txt':
            return self._load_with_docx2txt(file_path, metadata)
        elif self.loader_type == 'unstructured':
            return self._load_with_unstructured(file_path, metadata)
        else:  # python-docx (default)
            return self._load_with_python_docx(file_path, metadata)
    
    def _load_with_python_docx(self, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """Load using python-docx (default method)"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx is required for Word document processing. "
                "Install it with: pip install python-docx"
            )
        
        try:
            # Open Word document
            docx_doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraph_text = self._extract_text_from_paragraphs(docx_doc)
            
            # Extract text from tables
            table_text = self._extract_text_from_tables(docx_doc)
            
            # Combine all text
            full_text = paragraph_text
            if table_text:
                if full_text:
                    full_text += "\n\n" + table_text
                else:
                    full_text = table_text
            
            if not full_text.strip():
                logger.warning(f"Word document appears to be empty | File: {file_path}")
                raise ValueError("Word document contains no extractable text")
            
            # Extract document metadata
            doc_metadata = self._extract_metadata(docx_doc)
            
            # Add custom metadata if provided
            if metadata:
                doc_metadata.update(metadata)
            
            # Add file-specific metadata
            doc_metadata["source_file"] = file_path
            doc_metadata["file_type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            doc_metadata["char_count"] = len(full_text)
            doc_metadata["paragraph_count"] = len([p for p in docx_doc.paragraphs if p.text.strip()])
            doc_metadata["table_count"] = len(docx_doc.tables) if self.extract_tables else 0
            doc_metadata["loader_type"] = "python-docx"
            
            # Create Document object
            document = Document(page_content=full_text, metadata=doc_metadata)
            
            logger.info(
                f"Word document loaded successfully (python-docx) | "
                f"File: {file_path} | "
                f"Chars: {len(full_text)} | "
                f"Paragraphs: {doc_metadata['paragraph_count']} | "
                f"Tables: {doc_metadata['table_count']}"
            )
            
            return [document]
            
        except Exception as e:
            logger.error(f"Error loading Word file with python-docx | File: {file_path} | Error: {str(e)}", exc_info=True)
            raise RuntimeError(f"Failed to load Word file '{file_path}': {str(e)}")
    
    def _load_with_docx2txt(self, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """Load using Docx2txtLoader from langchain_community"""
        try:
            from langchain_community.document_loaders import Docx2txtLoader
        except ImportError:
            raise ImportError(
                "langchain-community is required for Docx2txtLoader. "
                "Install it with: pip install langchain-community"
            )
        
        try:
            loader = Docx2txtLoader(file_path)
            documents = loader.load()
            
            # Add custom metadata if provided
            if metadata:
                for doc in documents:
                    doc.metadata.update(metadata)
            
            # Add file-specific metadata
            for doc in documents:
                doc.metadata["source_file"] = file_path
                doc.metadata["file_type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                doc.metadata["loader_type"] = "docx2txt"
                if "char_count" not in doc.metadata:
                    doc.metadata["char_count"] = len(doc.page_content)
            
            logger.info(
                f"Word document loaded successfully (docx2txt) | "
                f"File: {file_path} | "
                f"Documents: {len(documents)}"
            )
            
            return documents
            
        except Exception as e:
            logger.error(f"Error loading Word file with docx2txt | File: {file_path} | Error: {str(e)}", exc_info=True)
            raise RuntimeError(f"Failed to load Word file '{file_path}': {str(e)}")
    
    def _load_with_unstructured(self, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """Load using UnstructuredWordDocumentLoader from langchain_community"""
        try:
            from langchain_community.document_loaders import UnstructuredWordDocumentLoader
        except ImportError:
            raise ImportError(
                "langchain-community is required for UnstructuredWordDocumentLoader. "
                "Install it with: pip install langchain-community"
            )
        
        try:
            # Use 'elements' mode for structured extraction
            loader = UnstructuredWordDocumentLoader(file_path, mode="elements")
            documents = loader.load()
            
            # Add custom metadata if provided
            if metadata:
                for doc in documents:
                    doc.metadata.update(metadata)
            
            # Add file-specific metadata
            for doc in documents:
                doc.metadata["source_file"] = file_path
                doc.metadata["file_type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                doc.metadata["loader_type"] = "unstructured"
                if "char_count" not in doc.metadata:
                    doc.metadata["char_count"] = len(doc.page_content)
            
            logger.info(
                f"Word document loaded successfully (unstructured) | "
                f"File: {file_path} | "
                f"Elements: {len(documents)} | "
                f"Element types: {set(doc.metadata.get('category', 'unknown') for doc in documents)}"
            )
            
            return documents
            
        except Exception as e:
            logger.error(f"Error loading Word file with unstructured | File: {file_path} | Error: {str(e)}", exc_info=True)
            raise RuntimeError(f"Failed to load Word file '{file_path}': {str(e)}")
    
    def load_bytes(self, docx_bytes: bytes, metadata: Optional[dict] = None) -> List[Document]:
        """
        Load Word document from bytes
        
        Args:
            docx_bytes: Word file content as bytes
            metadata: Optional metadata to add to documents
            
        Returns:
            List of Document objects
        """
        logger.debug(f"Loading Word document from bytes | Size: {len(docx_bytes)} bytes")
        
        # Validate Word magic bytes (ZIP file signature for .docx)
        if not docx_bytes.startswith(b'PK'):
            raise ValueError("Invalid Word file: File does not appear to be a valid .docx file")
        
        # Create a temporary file to save Word bytes
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(docx_bytes)
            temp_path = temp_file.name
        
        try:
            # Load from temporary file
            documents = self.load_file(temp_path, metadata)
            logger.info(f"Word document loaded from bytes | Documents: {len(documents)}")
            return documents
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_path)
            except Exception:
                pass
    
    def load_files(self, file_paths: List[str], metadata: Optional[dict] = None) -> List[Document]:
        """
        Load multiple Word files
        
        Args:
            file_paths: List of Word file paths
            metadata: Optional metadata to add to all documents
            
        Returns:
            List of Document objects from all files
        """
        all_documents = []
        for file_path in file_paths:
            documents = self.load_file(file_path, metadata)
            all_documents.extend(documents)
        return all_documents


# Convenience functions
def load_word_file(
    file_path: str,
    metadata: Optional[dict] = None,
    extract_tables: bool = True,
    extract_metadata: bool = True
) -> List[Document]:
    """
    Convenience function to load a Word file
    
    Args:
        file_path: Path to the Word file
        metadata: Optional metadata dictionary
        extract_tables: Whether to extract text from tables
        extract_metadata: Whether to extract document metadata
        
    Returns:
        List of Document objects
    """
    loader = WordDocumentLoader(
        extract_tables=extract_tables,
        extract_metadata=extract_metadata
    )
    return loader.load_file(file_path, metadata)


def load_word_bytes(
    docx_bytes: bytes,
    metadata: Optional[dict] = None,
    extract_tables: bool = True,
    extract_metadata: bool = True
) -> List[Document]:
    """
    Convenience function to load Word document from bytes
    
    Args:
        docx_bytes: Word file content as bytes
        metadata: Optional metadata dictionary
        extract_tables: Whether to extract text from tables
        extract_metadata: Whether to extract document metadata
        
    Returns:
        List of Document objects
    """
    loader = WordDocumentLoader(
        extract_tables=extract_tables,
        extract_metadata=extract_metadata
    )
    return loader.load_bytes(docx_bytes, metadata)
